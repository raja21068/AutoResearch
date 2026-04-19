"""
core/tools/file_reader.py — Read any file into text for LLM processing.

Supports: PDF, DOCX, Python, JS, Java, C/C++, Go, Rust, text, markdown,
          JSON, YAML, CSV, TSV, images, and more.
"""

import io
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)

# ── Extension maps ───────────────────────────────────────────

CODE_EXTS = {
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".c", ".cpp", ".h",
    ".hpp", ".go", ".rs", ".rb", ".php", ".swift", ".kt", ".scala",
    ".sh", ".bash", ".r", ".m", ".sql", ".html", ".css", ".vue",
    ".svelte", ".dart", ".lua", ".pl", ".ex", ".exs",
}
TEXT_EXTS = {
    ".txt", ".md", ".rst", ".log", ".env", ".gitignore", ".toml",
    ".ini", ".cfg", ".conf", ".xml", ".tex", ".bib", ".dockerfile",
}
DATA_EXTS = {".csv", ".tsv", ".json", ".yaml", ".yml", ".jsonl"}
IMG_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp", ".svg"}


# ── Public API ───────────────────────────────────────────────

def read_file(filepath: str) -> dict:
    """Read any file → {"filename", "content", "type", "pages"?, "rows"?}"""
    path = Path(filepath)
    if not path.exists():
        return {"filename": str(path), "content": "", "type": "missing", "error": "Not found"}

    ext = path.suffix.lower()
    name = path.name

    if ext == ".pdf":
        return _read_pdf(path, name)
    if ext in (".docx", ".doc"):
        return _read_docx(path, name)
    if ext in CODE_EXTS:
        return _read_text(path, name, "code")
    if ext in DATA_EXTS:
        return _read_dataset(path, name, ext)
    if ext in TEXT_EXTS:
        return _read_text(path, name, "text")
    if ext in IMG_EXTS:
        return {"filename": name, "type": "image",
                "content": f"[Image: {name}, {path.stat().st_size} bytes]"}
    try:
        return _read_text(path, name, "unknown")
    except Exception:
        return {"filename": name, "type": "binary",
                "content": f"[Binary: {name}, {path.stat().st_size} bytes]"}


def read_pdf_bytes(pdf_bytes: bytes, filename: str = "upload.pdf") -> dict:
    """Read PDF from in-memory bytes."""
    text, pages = "", 0
    # PyPDF2
    try:
        import PyPDF2
        r = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        pages = len(r.pages)
        for p in r.pages:
            t = p.extract_text()
            if t: text += t + "\n\n"
        if text.strip():
            return {"filename": filename, "content": text.strip(), "type": "pdf", "pages": pages}
    except Exception as e:
        logger.warning("PyPDF2: %s", e)
    # pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            pages = len(pdf.pages)
            for p in pdf.pages:
                t = p.extract_text()
                if t: text += t + "\n\n"
                for tbl in p.extract_tables():
                    for row in tbl:
                        text += " | ".join(str(c or "") for c in row) + "\n"
        if text.strip():
            return {"filename": filename, "content": text.strip(), "type": "pdf", "pages": pages}
    except Exception as e:
        logger.warning("pdfplumber: %s", e)
    return {"filename": filename, "content": "[Could not extract PDF text]", "type": "pdf", "pages": 0}


def read_docx_bytes(docx_bytes: bytes, filename: str = "upload.docx") -> dict:
    """Read DOCX from in-memory bytes."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(docx_bytes))
        text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text for cell in row.cells)
        return {"filename": filename, "content": text.strip(), "type": "docx", "pages": len(doc.paragraphs) // 30 + 1}
    except ImportError:
        return {"filename": filename, "content": "[python-docx not installed]", "type": "docx"}
    except Exception as e:
        return {"filename": filename, "content": f"[DOCX error: {e}]", "type": "docx"}


def read_dataset_bytes(data_bytes: bytes, filename: str, ext: str = ".csv") -> dict:
    """Read CSV/JSON/YAML dataset from bytes."""
    text = data_bytes.decode("utf-8", errors="ignore")
    return _parse_dataset(text, filename, ext)


def summarize_files(file_results: list[dict]) -> str:
    """Combine multiple file results into one context string."""
    parts = []
    for f in file_results:
        hdr = f"=== {f.get('filename', '?')} ({f.get('type', '?')}"
        if f.get("pages"): hdr += f", {f['pages']} pages"
        if f.get("rows"): hdr += f", {f['rows']} rows"
        hdr += ") ==="
        parts.append(f"{hdr}\n{f.get('content', '')}")
    return "\n\n".join(parts)


# ── Internal readers ─────────────────────────────────────────

def _read_pdf(path: Path, name: str) -> dict:
    return read_pdf_bytes(path.read_bytes(), name)

def _read_docx(path: Path, name: str) -> dict:
    return read_docx_bytes(path.read_bytes(), name)

def _read_text(path: Path, name: str, ftype: str) -> dict:
    content = path.read_text(encoding="utf-8", errors="ignore")
    if len(content) > 100_000:
        content = content[:100_000] + f"\n[...truncated, {len(content)} total chars...]"
    return {"filename": name, "content": content, "type": ftype}

def _read_dataset(path: Path, name: str, ext: str) -> dict:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return _parse_dataset(text, name, ext)

def _parse_dataset(text: str, name: str, ext: str) -> dict:
    rows = 0
    preview = ""
    if ext in (".csv", ".tsv"):
        lines = text.strip().split("\n")
        rows = len(lines) - 1  # minus header
        # Show header + first 20 rows + stats
        preview_lines = lines[:21]
        preview = "\n".join(preview_lines)
        if rows > 20:
            preview += f"\n... ({rows} total rows)"
        # Column info
        if lines:
            cols = lines[0].split("," if ext == ".csv" else "\t")
            preview = f"Columns ({len(cols)}): {', '.join(cols[:20])}\n\n{preview}"
    elif ext in (".json", ".jsonl"):
        import json as _json
        try:
            data = _json.loads(text)
            if isinstance(data, list):
                rows = len(data)
                preview = _json.dumps(data[:5], indent=2)
                if rows > 5:
                    preview += f"\n... ({rows} total items)"
            elif isinstance(data, dict):
                rows = len(data)
                preview = _json.dumps(data, indent=2)[:3000]
            else:
                preview = text[:3000]
        except _json.JSONDecodeError:
            # JSONL
            lines = text.strip().split("\n")
            rows = len(lines)
            preview = "\n".join(lines[:10])
            if rows > 10:
                preview += f"\n... ({rows} total lines)"
    elif ext in (".yaml", ".yml"):
        preview = text[:3000]
        rows = text.count("\n")
    else:
        preview = text[:3000]

    return {"filename": name, "content": preview, "type": "dataset", "rows": rows}
